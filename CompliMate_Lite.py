import os
import json
import faiss
import hashlib
import pickle
from docx import Document
from pdf2docx import Converter
from sentence_transformers import SentenceTransformer
from difflib import SequenceMatcher

# ------------------------
# Tunable thresholds / constants
# ------------------------
HEADING_FUZZY_THRESHOLD = 0.8
KEYWORD_SIM_THRESHOLD = 0.6

# ------------------------
# Role-Keyword Mapping
# ------------------------

ROLE_FILE_MAP = {
    "retail": [
        "FORM XIV - Storage in Tanks for Pump Outfits", "FORM XVIII - Decanting Kerosene",
        "FORM XIX - On-Site Refuelling by Tanker", "Approval for Fabrication Shop & Safety Fittings",
        "Fuel Dispenser Compliance", "Retail pump installation layout", "Pressure Vacuum Valve",
        "Emergency Shut-off Valve", "Fusible Link", "Retail Transport Safety Distance",
        "TPIA (Third Party Inspection Agency)", "Competent Person", "Storage in Non-Bulk Containers",
        "Schedule 7 / Schedule 8", "Licensing of retail outlets", "HAZOP / QRA for retail outlets",
        "Safety fittings approval for bowsers/tankers", "Statutory Clearance for pump installations",
        "Public Safety", "Retail licensing limits (Petroleum Class A/B/C)",
        "Transport by road in tankers (retail refueling)", "National Single Window Portal (NSWP)",
        "Offences and Penalties related to retail storage"
    ],
    "non_retail": [
        "FORM XV - Import & Storage in Installation", "FORM XVI - Import & Storage Otherwise Than in Bulk",
        "FORM XI - Carriage by Land (Mechanically Propelled Vehicles)", "ISO Tank Container Permissions",
        "Pipeline Approvals", "Port / Jetties Approval", "Refinery Approvals",
        "Ex Electrical Apparatus Approval", "Emergency Vent", "Spark Arrester", "Storage in Bulk",
        "Flash Point classification", "Import permission", "Port NOC", "MoEF Clearance",
        "EIA / HAZOP / QRA Reports", "Corrosion Protection", "Hazardous Area Classification (Zone 0,1,2)",
        "Refinery layout and commissioning", "Statutory Clearance for refineries/jetties",
        "Fabrication Shop for Tank Trucks", "Approval for Ex equipment (import/manufacture)",
        "Rules by Central Government", "Power of Inspection", "Testing of Petroleum", "Delegation of powers"
    ]
}

shared_items = [
    "Definition of Petroleum", "Petroleum Class A/B/C", "Licensing Authority", "Competent Person",
    "TPIA", "Fire Fighting Facilities", "Emergency Shut-off Valve", "National Single Window Portal (NSWP)",
    "Public Safety", "Revocation and Renewal of License", "Hazardous Area Classification",
    "Safety Distance Requirements", "Offences and Penalties", "Inspection and Testing",
    "Storage Regulations", "Licensing of Petroleum Activities", "Power to Make Rules", "Delegation of Powers"
]

for item in shared_items:
    for role in ROLE_FILE_MAP:
        if item not in ROLE_FILE_MAP[role]:
            ROLE_FILE_MAP[role].append(item)


# ------------------------
# CompliMate_lite Class
# ------------------------

class CompliMateLite:
   import os
import json
import faiss
from sentence_transformers import SentenceTransformer

class CompliMateLite:
    def __init__(self,
                 rag_folder_lite=None,   # Path to RAG docs folder
                 meta_folder_lite=None): # Path to META folder
        
        base_dir = os.path.dirname(__file__)  # Where this file is located

        # Default to folders inside repo if not provided
        if rag_folder_lite is None:
            rag_folder_lite = os.path.join(base_dir, "RAG_folder_lite")
        if meta_folder_lite is None:
            meta_folder_lite = os.path.join(base_dir, "meta_folder_lite")

        self.rag_folder_lite = rag_folder_lite
        self.meta_folder_lite = meta_folder_lite
        os.makedirs(self.meta_folder_lite, exist_ok=True)

        # File paths inside META folder
        self.meta_file = os.path.join(self.meta_folder_lite, "section_meta_lite.json")
        self.processed_file = os.path.join(self.meta_folder_lite, "processed_lite.json")
        self.index_file = os.path.join(self.meta_folder_lite, "section_index_lite.faiss")
        self.paragraph_index_file = os.path.join(self.meta_folder_lite, "paragraph_index_lite.faiss")
        self.paragraphs_file = os.path.join(self.meta_folder_lite, "paragraphs_lite.pkl")

        # Persistence for learned roles/shared headings
        self.roles_file = os.path.join(self.meta_folder_lite, "roles_map.json")
        self.shared_file = os.path.join(self.meta_folder_lite, "shared_items.json")

        # Model setup
        self.model = SentenceTransformer("multi-qa-MiniLM-L6-cos-v1")
        dim = self.model.get_sentence_embedding_dimension()
        self.index = None
        self.paragraph_index = None

        # In-memory stores
        self.meta = []
        self.paragraphs = []
        self.processed = {}

        # Load processed map early if present
        if os.path.exists(self.processed_file):
            try:
                with open(self.processed_file, "r", encoding="utf-8") as f:
                    self.processed = json.load(f)
            except Exception as e:
                print(f"[WARN] Could not load processed file: {e}")

        # Merge persisted roles/shared
        self._load_persisted_roles()

        # If all index/meta files exist → load, else build fresh
        required = [self.meta_file, self.processed_file,
                    self.index_file, self.paragraph_index_file, self.paragraphs_file]
        if all(os.path.exists(p) for p in required):
            self._load_all()
            print("Meta entries:", len(self.meta))
            print("Section index size:", self.index.ntotal)
            print("Paragraphs entries:", len(self.paragraphs))
            print("Paragraph index size:", self.paragraph_index.ntotal)
        else:
            self.index = faiss.IndexFlatL2(dim)
            self.paragraph_index = faiss.IndexFlatL2(dim)
            self._build_index()


    def _get_file_hash(self, filepath):
        hasher = hashlib.md5()
        with open(filepath, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    # ------------------------
    # Persistence helpers
    # ------------------------
    def _atomic_write_json(self, path, data):
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
            f.flush()
            os.fsync(f.fileno())
        os.replace(tmp, path)

    def _atomic_write_pickle(self, path, data):
        tmp = path + ".tmp"
        with open(tmp, "wb") as f:
            pickle.dump(data, f)
            f.flush()
            os.fsync(f.fileno())
        os.replace(tmp, path)

    def _atomic_write_faiss(self, index_obj, path):
        tmp = path + ".tmp"
        faiss.write_index(index_obj, tmp)
        os.replace(tmp, path)

    # ------------------------
    # Roles / shared persistence
    # ------------------------
    def _load_persisted_roles(self):
        try:
        # roles_map.json
            if os.path.exists(self.roles_file):
                file_hash = self._get_file_hash(self.roles_file)
                if self.processed.get("roles_file") != file_hash:
                    with open(self.roles_file, "r", encoding="utf-8") as f:
                        persisted = json.load(f)
                    for role, kws in persisted.items():
                        if role not in ROLE_FILE_MAP:
                            ROLE_FILE_MAP[role] = []
                        for kw in kws:
                            if kw not in ROLE_FILE_MAP[role]:
                                ROLE_FILE_MAP[role].append(kw)
                    self.processed["roles_file"] = file_hash

        # shared_items.json
            if os.path.exists(self.shared_file):
                file_hash = self._get_file_hash(self.shared_file)
                if self.processed.get("shared_file") != file_hash:
                    with open(self.shared_file, "r", encoding="utf-8") as f:
                        persisted_shared = json.load(f)
                    for it in persisted_shared:
                        if it not in shared_items:
                            shared_items.append(it)
                            for role in ROLE_FILE_MAP:
                                if it not in ROLE_FILE_MAP[role]:
                                    ROLE_FILE_MAP[role].append(it)
                self.processed["shared_file"] = file_hash
        except Exception as e:
            print(f"[WARN] Failed to load persisted roles/shared: {e}")

    def _persist_roles_and_shared(self):
        try:
        # persist ROLE_FILE_MAP (only lists)
            self._atomic_write_json(self.roles_file, ROLE_FILE_MAP)
            self.processed["roles_file"] = self._get_file_hash(self.roles_file)
        except Exception as e:
            print(f"[WARN] Failed to persist ROLE_FILE_MAP: {e}")

        try:
            self._atomic_write_json(self.shared_file, shared_items)
            self.processed["shared_file"] = self._get_file_hash(self.shared_file)
        except Exception as e:
            print(f"[WARN] Failed to persist shared_items: {e}")

    # persist updated processed map
        try:
            self._atomic_write_json(self.processed_file, self.processed)
        except Exception as e:
            print(f"[WARN] Failed to persist updated processed file after roles/shared: {e}")


    # ------------------------
    # Loading & Building
    # ------------------------
    def _load_all(self):
        # load meta and processed
        with open(self.meta_file, "r", encoding="utf-8") as f:
            self.meta = json.load(f)
        with open(self.processed_file, "r", encoding="utf-8") as f:
            self.processed = json.load(f)

        # load paragraph metadata
        with open(self.paragraphs_file, "rb") as f:
            self.paragraphs = pickle.load(f)

        # load FAISS indexes (guard against corruption)
        dim = self.model.get_sentence_embedding_dimension()
        try:
            self.index = faiss.read_index(self.index_file)
        except Exception as e:
            print(f"[WARN] Failed to load section index (will rebuild in-memory): {e}")
            self.index = faiss.IndexFlatL2(dim)
            if self.meta:
                sections = [entry['content'] for entry in self.meta]
                if sections:
                    vecs = self.model.encode(sections, convert_to_numpy=True).astype("float32")
                    self.index.add(vecs)
            # persist repaired index
            try:
                self._atomic_write_faiss(self.index, self.index_file)
            except Exception as e2:
                print(f"[WARN] Could not persist repaired section index: {e2}")

        try:
            self.paragraph_index = faiss.read_index(self.paragraph_index_file)
        except Exception as e:
            print(f"[WARN] Failed to load paragraph index (will rebuild in-memory): {e}")
            self.paragraph_index = faiss.IndexFlatL2(dim)
            if self.paragraphs:
                para_texts = [p['paragraph'] for p in self.paragraphs]
                if para_texts:
                    vecs = self.model.encode(para_texts, convert_to_numpy=True).astype("float32")
                    self.paragraph_index.add(vecs)
            try:
                self._atomic_write_faiss(self.paragraph_index, self.paragraph_index_file)
            except Exception as e2:
                print(f"[WARN] Could not persist repaired paragraph index: {e2}")
            
            


    def _build_index(self):
        # reset
        self.meta = []
        self.paragraphs = []
        # preserve processed if present else reset
        if not os.path.exists(self.processed_file):
            self.processed = {}

        # collect data
        meta, sections, paragraphs = self._load_documents()

        if not sections:
            print("⚠️ No documents found for indexing.")
            return

        # ensure fresh indexes
        dim = self.model.get_sentence_embedding_dimension()
        self.index = faiss.IndexFlatIP(dim)
        self.paragraph_index = faiss.IndexFlatIP(dim)

        # encode & add sections (guarded)
        if sections:
            section_vectors = self.model.encode(sections, convert_to_numpy=True).astype("float32")
            faiss.normalize_L2(section_vectors)
            self.index.add(section_vectors)

        # encode & add paragraphs (use paragraph text) (guarded)
        para_texts = [p['paragraph'] for p in paragraphs]
        if para_texts:
            para_vectors = self.model.encode(para_texts, convert_to_numpy=True).astype("float32")
            faiss.normalize_L2(para_vectors)
            self.paragraph_index.add(para_vectors)

        # persist meta/paragraphs/processed/indexes
        self.meta = meta
        self.paragraphs = paragraphs

        try:
            self._atomic_write_json(self.meta_file, self.meta)
        except Exception as e:
            print(f"[WARN] Failed to write meta file: {e}")
        try:
            self._atomic_write_pickle(self.paragraphs_file, self.paragraphs)
        except Exception as e:
            print(f"[WARN] Failed to write paragraphs file: {e}")
        try:
            self._atomic_write_json(self.processed_file, self.processed)
        except Exception as e:
            print(f"[WARN] Failed to write processed file: {e}")

        try:
            self._atomic_write_faiss(self.index, self.index_file)
        except Exception as e:
            print(f"[WARN] Failed to persist section index: {e}")
        try:
            self._atomic_write_faiss(self.paragraph_index, self.paragraph_index_file)
        except Exception as e:
            print(f"[WARN] Failed to persist paragraph index: {e}")

        print(f"✅ Index built: {len(sections)} sections, {len(paragraphs)} paragraphs.")

    # ------------------------
    # Document loading & extraction
    # ------------------------
    def _load_documents(self):
        meta = []
        sections = []
        paragraphs = []

        # self.processed is expected to exist (either loaded or reset)
        for filename in os.listdir(self.rag_folder_lite):
            lower = filename.lower()
            fullpath = os.path.join(self.rag_folder_lite, filename)
            if lower.endswith(".pdf"):
                docx_path = self._convert_pdf_to_docx(fullpath)
                if not docx_path:
                    print(f"[WARN] PDF -> DOCX conversion failed for {filename}, skipping.")
                    continue
            elif lower.endswith(".docx"):
                docx_path = fullpath
            else:
                continue

            # hash source file (pdf or docx) and use relpath key
            source_path = fullpath
            try:
                file_hash = self._get_file_hash(source_path)
            except Exception as e:
                print(f"[ERROR] Failed hashing {source_path}: {e}")
                continue    

            key = os.path.relpath(source_path, self.rag_folder_lite)
            if self.processed.get(key) == file_hash:
                continue  # unchanged

           

            # extract sections & paragraphs
            section_data, paragraph_data = self._extract_sections(docx_path)
            meta.extend(section_data)
            sections.extend([s['content'] for s in section_data])
            paragraphs.extend(paragraph_data)

            # update processed map
            self.processed[key] = file_hash

        # update meta file AFTER processing (will be saved by _build_index)
        return meta, sections, paragraphs

    def _extract_sections(self, docx_path):
        section_data = []
        paragraph_data = []

        sections = self._split_into_sections(docx_path)
        for heading, content in sections:
            access_tag = self._infer_access_tag(heading)

            section_record = {
                "filename": os.path.basename(docx_path),
                "heading": heading,
                "content": content,
                "access_tag": access_tag,
            }
            section_data.append(section_record)

            # paragraph-level: keep all paragraphs (no length filtering)
            for para in content.split("\n"):
                p = para.strip()
                if p:
                    paragraph_data.append({
                        "filename": os.path.basename(docx_path),
                        "section_heading": heading,
                        "section_content": content,
                        "paragraph": p,
                        "access_tag": access_tag
                    })

        return section_data, paragraph_data

    def _split_into_sections(self, filepath):
        try:
            doc = Document(filepath)
            sections = []
            current_heading = "Introduction"
            current_content = []
            table_index = 0
            elements = list(doc.element.body.iterchildren())

            for el in elements:
                if el.tag.endswith("}p"):
                    para = next((p for p in doc.paragraphs if p._element == el), None)
                    if not para:
                        continue
                    if self._is_heading(para):
                        if current_content:
                            sections.append((current_heading, "\n".join(current_content)))
                        current_heading = para.text.strip()
                        current_content = []
                    else:
                        current_content.append(para.text.strip())
                elif el.tag.endswith("}tbl"):
                    table = next((t for t in doc.tables if t._element == el), None)
                    if table:
                        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                        current_content.append("\n".join(rows))

            if current_content:
                sections.append((current_heading, "\n".join(current_content)))
            return sections
        except Exception as e:
            print(f"[ERROR] Failed to parse {filepath}: {e}")
            return []

    def _is_heading(self, para):
        text = para.text.strip()
        if not text:
            return False
        try:
            # Bold run OR style name containing "heading"
            if any(run.bold and (run.font.size is None or run.font.size.pt >= 10) for run in para.runs):
                return True
            if para.style and getattr(para.style, "name", None) and "heading" in para.style.name.lower():
                return True
        except Exception:
            # defensive
            pass
        return False

    # ------------------------
    # Role / access helpers
    # ------------------------
    def _keyword_similarity(self, a, b):
        set_a = set(a.lower().split())
        set_b = set(b.lower().split())
        if not set_a or not set_b:
            return 0.0
        return len(set_a & set_b) / len(set_a | set_b)

    def _infer_access_tag(self, heading):
        max_sim = 0.0
        best_role = None
        for role, keywords in ROLE_FILE_MAP.items():
            for keyword in keywords:
                sim = self._keyword_similarity(heading, keyword)
                if sim > max_sim:
                    max_sim = sim
                    best_role = role

        # threshold tuned lower to be inclusive; unmatched default -> shared
        if max_sim >= KEYWORD_SIM_THRESHOLD:
            return best_role
        else:
            if heading not in shared_items:
                shared_items.append(heading)
                for role in ROLE_FILE_MAP:
                    if heading not in ROLE_FILE_MAP[role]:
                        ROLE_FILE_MAP[role].append(heading)
                # persist learned roles/shared
                try:
                    self._persist_roles_and_shared()
                except Exception as e:
                    print(f"[WARN] Failed persisting learned roles/shared: {e}")
            return "shared"

    def _is_authorized(self, user_role, heading):
        # allow shared headings for all roles
        meta_entry = next((m for m in self.meta if m.get("heading") == heading), None)
        if meta_entry and meta_entry.get("access_tag") == "shared":
            return True
        for kw in ROLE_FILE_MAP.get(user_role, []):
            if self._keyword_similarity(heading, kw) >= KEYWORD_SIM_THRESHOLD:
                return True
        return False

    # ------------------------
    # Retrieval
    # ------------------------
    def _heading_fuzzy_matches(self, query, threshold=HEADING_FUZZY_THRESHOLD):
        """Return list of meta entries where heading fuzzy/substring matches query.
           Sorted by ratio desc."""
        if not self.meta:
            return []
        q = query.strip().lower()
        matches = []
        for entry in self.meta:
            heading = entry.get("heading", "")
            h_lower = heading.lower()
            ratio = SequenceMatcher(None, q, h_lower).ratio()
            if q in h_lower or ratio >= threshold:
                matches.append((ratio, entry))
        matches.sort(key=lambda x: x[0], reverse=True)
        return [e for _, e in matches]

    def _semantic_paragraph_search(self, prompt, top_k=5):
        if not self.paragraph_index or not self.paragraphs or getattr(self.paragraph_index, "ntotal", 0) == 0:
            return []
        qvec = self.model.encode([prompt], convert_to_numpy=True).astype("float32")
        faiss.normalize_L2(qvec)
        D, I = self.paragraph_index.search(qvec, top_k)

        results = []
        for idx in I[0]:
            if idx < 0 or idx >= len(self.paragraphs):
                continue
            results.append(self.paragraphs[idx])
        return results

    def query(self, prompt, user_role, top_k=5):
        # Phase 1: Fuzzy heading matches
        heading_matches = self._heading_fuzzy_matches(prompt, threshold=HEADING_FUZZY_THRESHOLD)
        fuzzy_results = []
        seen_sections = set()

        for entry in heading_matches:
            sec_id = (entry['filename'], entry['heading'])
            if sec_id not in seen_sections and self._is_authorized(user_role, entry['heading']):
                fuzzy_results.append({
                    'filename': entry['filename'],
                    'section_heading': entry['heading'],
                    'section_content': entry['content'],
                    'access_tag': entry.get('access_tag', 'shared')
                })
                seen_sections.add(sec_id)

        # Phase 2: Semantic matches
        sem_results = self._semantic_paragraph_search(prompt, top_k=top_k)
        semantic_results = []
        for para in sem_results:
            sec_id = (para['filename'], para['section_heading'])
            if sec_id not in seen_sections and self._is_authorized(user_role, para['section_heading']):
                semantic_results.append({
                    'filename': para['filename'],
                    'section_heading': para['section_heading'],
                    'section_content': para['section_content'],
                    'access_tag': para.get('access_tag', 'shared')
                })
                seen_sections.add(sec_id)

        if not fuzzy_results and not semantic_results:
            return "No relevant sections found."

        output_parts = []
        if fuzzy_results:
            output_parts.append(self._format_response(fuzzy_results))
        if semantic_results:
            output_parts.append(self._format_response(semantic_results))

        return "\n\n".join(output_parts)


    def _format_response(self, results):
        output = []
        for res in results:
            output.append(
                f"**File:** {res['filename']}\n\n"
                f"**Heading:** {res['section_heading']}\n\n"
                f"**Content**\n{res['section_content']}\n\n")
        

        return "\n\n".join(output).strip()

    # ------------------------
    # Utilities
    # ------------------------
    def _convert_pdf_to_docx(self, pdf_path):
        # safer replacement for '.pdf' suffix
        if pdf_path.lower().endswith(".pdf"):
            docx_path = pdf_path[:-4] + ".docx"
        else:
            docx_path = pdf_path + ".docx"
        try:
            if not os.path.exists(docx_path):
                cv = Converter(pdf_path)
                cv.convert(docx_path)
                cv.close()
        except Exception as e:
            print(f"[ERROR] PDF conversion failed for {pdf_path}: {e}")
            return None

        return docx_path
